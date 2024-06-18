import streamlit as st
import json
from docx import Document
from io import BytesIO

# Function to create Word document from user profile
def create_word_doc(user_profile):
    doc = Document()
    doc.add_heading('User Profile', 0)

    def add_dict_to_doc(d, doc, level=0):
        for key, value in d.items():
            if isinstance(value, dict):
                doc.add_heading(key, level=level + 1)
                add_dict_to_doc(value, doc, level + 1)
            elif isinstance(value, list):
                doc.add_heading(key, level=level + 1)
                for item in value:
                    if isinstance(item, dict):
                        add_dict_to_doc(item, doc, level + 2)
                    else:
                        doc.add_paragraph(str(item))
            else:
                doc.add_paragraph(f"{key}: {value}")

    add_dict_to_doc(user_profile, doc)
    
    return doc

# Title and Introduction
st.title("Foodeasy - Personalized Meal Planning Assistant")
st.markdown("""
Hey friend! 👋 Hope your day is as awesome as your favorite meal! We're here to make your mealtime even more fun and stress-free! 🍽️✨

Let's chat about your eating habits and preferences so we can create something perfect for you! Ready?
""")

# File upload or text input for food recommendations
st.markdown("**Got your microbiome or food sensitivity test results? Awesome!**")
uploaded_file = st.file_uploader("Upload your food recommendations file", type=["txt", "pdf", "docx"])
results_input = st.text_area("Or type in your results here")

# Number of people for meal plan
st.markdown("**Awesome! Let's go! 🥳**")
st.markdown("**How many people are we preparing the meal plan for? 👨🍳 Just you, or a whole group? 👥**")
meal_plan_for = st.radio("Select an option:", ["Just me", "A group"])

# Top health goals
st.markdown("**What are your top health goals with Foodeasy? 🍽️ Pick up to 3 that matter most to you:**")
health_goals = st.multiselect(
    "Select your top health goals (up to 3):",
    [
        "Lose weight 🏋️‍♂️", "Save money 💸", "Simplify cooking 👩🍳", "Save time ⏰",
        "Try new things 🌟", "Improve health 💪", "Grocery shop less 🛒", "Waste less food 🌱", "Other (let us know!) 📝"
    ],
    max_selections=3
)
if "Other (let us know!) 📝" in health_goals:
    other_goal = st.text_input("Please specify your other goal:")

# Overview of usual meals
st.markdown("**Can you give me a quick overview of what you usually eat for each meal?**")
breakfast = st.text_area("Breakfast: 🍳", key="breakfast")
lunch = st.text_area("Lunch: 🥪", key="lunch")
dinner = st.text_area("Dinner: 🍝", key="dinner")
snacks = st.text_area("Snacks: 🍏", key="snacks")

# Additional dietary preferences
st.markdown("**How would you describe your eating style? 🍽️**")
eating_style = st.radio(
    "Select your eating style:",
    ["Omnivore", "Pescatarian", "Vegetarian", "Vegan", "No restrictions", "Other (please tell us more!)"]
)
if eating_style == "Other (please tell us more!)":
    other_eating_style = st.text_input("Please specify your eating style:")

st.markdown("**Do you have any other dietary needs? 🍽️ You can always update this later.**")
dietary_needs = st.multiselect(
    "Select any dietary needs:",
    ["Dairy-Free 🧀🚫", "Gluten-Free 🌾🚫", "Soy-Free 🌱🚫", "Tree Nut-Free 🌰🚫", "Peanut-Free 🥜🚫", "Egg-Free 🥚🚫", "Shellfish-Free 🦐🚫"]
)

other_restrictions = st.text_area("Do you have any other food restrictions or allergies that I might have missed? 📝")

# Nutrition preferences
st.markdown("**Do you have any nutrition preferences? 🥗 Your choices help us suggest recipes that fit how you want to eat. Pick any that apply:**")
nutrition_preferences = st.multiselect(
    "Select your nutrition preferences:",
    ["Carb conscious (<35g) 🍞🚫", "High protein (>25g) 💪", "Gut friendly 🌿", "Less sodium (<480mg) 🧂🚫", "Anti-inflammatory 🌟", 
     "Immunity boosting 🛡️", "Less sugar (<2g added sugar) 🍬🚫", "Nothing specific 🤷‍♂️"]
)

# Meat preferences
st.markdown("**In what order do you like these types of meat? 🥩🍗 Just let us know your preferences!**")
meat_preferences = st.multiselect(
    "Select your preferred types of meat:",
    ["Beef + Bison (ex: meatballs, ground beef)", "Poultry (ex: grilled chicken, ground turkey)", "Pork (ex: carnitas, sliced ham)", 
     "Lamb (ex: lamb cubes, lamb chops)", "None of them 🚫", "Doesn’t matter, I like them all! 😋"]
)

# Seafood preferences
st.markdown("**What kinds of seafood do you like? 🐟🦐 Let us know your favorites in order!**")
seafood_preferences = st.multiselect(
    "Select your preferred types of seafood:",
    ["Fresh Fish (ex: cod, salmon)", "Smoked Fish (ex: lox, hot smoked salmon)", "Shellfish (ex: shrimp, crab, lobster)", 
     "Mollusks (ex: clams, oysters, mussels)", "Squid and Octopus", "None of them 🚫", "Doesn’t matter, I like them all! 😋"]
)

# Plant-based protein preferences
st.markdown("**What kinds of plant-based proteins do you like? 🌱🍲 Let us know your favorites in order!**")
plant_protein_preferences = st.multiselect(
    "Select your preferred plant-based proteins:",
    ["Tofu + Tempeh (ex: tofu nuggets, yuba noodles)", "Beans + Lentils (ex: simmered black beans, lentil dal)", 
     "‘Meat’ alternatives (ex: Beyond Meat, plant-based chorizo)", "None of them 🚫", "Doesn’t matter, I like them all! 😋"]
)

# Foods disliked or avoided
disliked_foods = st.text_area("Are there any foods you don't eat or don't like? 🍅🥔 (For example: tomatoes, eggplants, mushrooms, cilantro, bell pepper, potato, garlic, onions, or anything else?)")

# Meal and snack preferences
st.markdown("**What sounds good for breakfast? 🌞🍽️**")
breakfast_preferences = st.multiselect(
    "Select your breakfast preferences:",
    ["Sandwiches + wraps 🥪", "Smoothies 🍓", "Baked goods 🥐", "Cereal 🥣", "Coffee + tea ☕", "Oatmeal + granola 🍯", 
     "Yogurt 🥄", "Eggs 🍳", "Other"]
)

st.markdown("**What sounds good for lunch and dinner? 🌯🍔**")
lunch_dinner_preferences = st.multiselect(
    "Select your lunch and dinner preferences:",
    ["Sandwiches 🥪", "Bowls 🍲", "Wraps 🌯", "Main + Sides 🍛", "Stir-Fries 🍜", "Pastas 🍝", "Pizzas 🍕", "Salads 🥗", 
     "Burgers 🍔", "Tacos 🌮", "Other"]
)

st.markdown("**What snacks look good? 🍿🍏**")
snack_preferences = st.multiselect(
    "Select your snack preferences:",
    ["Chips, puffs + popcorn 🍿", "Juices 🥤", "Crackers + pretzels 🥨", "Dried fruit 🍇", "Snack packs 🎒", "Smoothies 🍓", 
     "Fresh fruit + veggies 🍎🥕", "Mini meals 🍱", "Cheese 🧀", "Bars + bites 🍫", "Jerky 🍖", "Pickles + olives 🥒🫒", 
     "Nuts, seeds + trail mix 🌰", "Other"]
)

# Cooking and meal preparation habits
st.markdown("**Let's talk about your meals and how much time you spend cooking or preparing them. ⏲️**")
breakfast_time = st.number_input("Do you usually have breakfast? 🍳 How much time do you spend making it? (minutes)", min_value=0, step=1)
morning_snack_time = st.number_input("Do you usually have a snack between breakfast and lunch? 🍏 How much time do you spend making it? (minutes)", min_value=0, step=1)
lunch_time = st.number_input("Do you usually have lunch? 🥪 How much time do you spend making it? (minutes)", min_value=0, step=1)
afternoon_snack_time = st.number_input("Do you usually have a snack between lunch and dinner? 🍎 How much time do you spend making it? (minutes)", min_value=0, step=1)
dinner_time = st.number_input("Do you usually have dinner? 🍝 How much time do you spend making it? (minutes)", min_value=0, step=1)

# Cooking frequency
st.markdown("**How often do you cook at home? 👩🍳**")
cooking_frequency = st.radio(
    "Select your cooking frequency:",
    ["Every day", "4 to 5 times per week", "2 to 3 times per week", "Once per week"]
)

# Preferences for leftovers and repeating recipes
st.markdown("**How do you feel about leftovers and repeating recipes? 🍲**")
leftovers_preference = st.radio(
    "Select your preference for leftovers:",
    ["Love them!", "Like them", "Neutral", "Dislike them", "Hate them"]
)

repeating_recipes_preference = st.radio(
    "How do you feel about repeating recipes?",
    ["Love it!", "Like it", "Neutral", "Dislike it", "Hate it"]
)

# Kitchen appliances
st.markdown("**Which kitchen appliances do you have? 🍽️**")
kitchen_appliances = st.multiselect(
    "Select your kitchen appliances:",
    ["Oven", "Stove", "Microwave", "Toaster", "Blender", "Food Processor", "Slow Cooker", "Instant Pot", "Air Fryer", "Grill", "None"]
)

# Grocery shopping preferences
st.markdown("**How do you like to shop for groceries? 🛒**")
grocery_shopping_method = st.radio(
    "Select your grocery shopping method:",
    ["In-store", "Online", "Both"]
)

grocery_stores = st.multiselect(
    "Select your preferred grocery stores:",
    ["Local farmer's market", "Whole Foods", "Trader Joe's", "Costco", "Walmart", "Target", "Kroger", "Other"]
)

non_preferred_stores = st.text_area("Are there any stores or brands you prefer not to use?")

# Delivery instructions and contact information
st.markdown("**Any delivery instructions or contact information we should know about? 📦**")
delivery_instructions = st.text_area("Please specify any delivery instructions:")

contact_name = st.text_input("Name:")
contact_email = st.text_input("Email:")
contact_phone = st.text_input("Phone:")
contact_address = st.text_input("Delivery Address:")

# Create user profile
if st.button("Submit"):
    user_profile = {
        "health_goals": health_goals + [other_goal] if "Other (let us know!) 📝" in health_goals else health_goals,
        "meal_history_overview": {
            "breakfast": breakfast.split(', '),
            "lunch": lunch.split(', '),
            "dinner": dinner.split(', '),
            "snacks": snacks.split(', ')
        },
        "eating_style": other_eating_style if eating_style == "Other (please tell us more!)" else eating_style,
        "dietary_needs": dietary_needs,
        "other_food_restrictions": other_restrictions,
        "nutrition_preferences": nutrition_preferences,
        "meat_order_preference": meat_preferences,
        "seafood_order_preference": seafood_preferences,
        "plant_based_protein_order_preference": plant_protein_preferences,
        "foods_not_eaten": disliked_foods,
        "breakfast_preferences": breakfast_preferences,
        "lunch_dinner_preferences": lunch_dinner_preferences,
        "snack_preferences": snack_preferences,
        "meal_information": [
            {"breakfast": bool(breakfast_time), "breakfast_preparation_time": f"{breakfast_time} minutes"},
            {"breakfast_to_lunch_snack": bool(morning_snack_time), "breakfast_to_lunch_snack_preparation_time": f"{morning_snack_time} minutes"},
            {"lunch": bool(lunch_time), "lunch_preparation_time": f"{lunch_time} minutes"},
            {"lunch_to_dinner_snack": bool(afternoon_snack_time), "lunch_to_dinner_snack_preparation_time": f"{afternoon_snack_time} minutes"},
            {"dinner": bool(dinner_time), "dinner_preparation_time": f"{dinner_time} minutes"}
        ],
        "cooking_frequency": cooking_frequency,
        "leftovers_and_recipe_repetition": {
            "repeating_recipes": repeating_recipes_preference,
            "leftovers": leftovers_preference
        },
        "kitchen_appliances": kitchen_appliances,
        "grocery_shopping_method": grocery_shopping_method,
        "grocery_stores": grocery_stores,
        "non_preferred_stores_or_brands": non_preferred_stores,
        "delivery_instructions": delivery_instructions,
        "contact_info": {
            "name": contact_name,
            "email": contact_email,
            "phone": contact_phone,
            "delivery_address": contact_address
        }
    }
    
    # Create the Word document
    doc = create_word_doc(user_profile)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("Your profile has been created! 🎉")
    
    # Create a download button for the Word document
    st.download_button(
        label="Download Profile as Word Document",
        data=buffer,
        file_name="user_profile.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
